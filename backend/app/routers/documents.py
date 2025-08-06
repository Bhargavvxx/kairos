# backend/app/routers/documents.py

import logging
import json
import io
import uuid
import hashlib
from pathlib import Path
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Set, Tuple
import asyncio
from concurrent.futures import ThreadPoolExecutor

from fastapi import (
    APIRouter, UploadFile, File, HTTPException, Depends, 
    BackgroundTasks, Form, Request, status
)
from fastapi.responses import JSONResponse
from sse_starlette.sse import EventSourceResponse
import chardet
import magic
import aiofiles
from pydantic import BaseModel, Field

# Document processing libraries
try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Import services and schemas
from app.dependencies import (
    get_vector_service,
    get_graph_service,
    get_llm_client,
)
from app.services.vector_service import VectorService
from app.services.graph_service import GraphService
from app.models.schemas import DocumentMetadata, ProcessingStatus

# --- Configuration ---

class DocumentConfig:
    """Configuration for document processing."""
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB default
    LARGE_FILE_THRESHOLD = 5 * 1024 * 1024  # 5MB for streaming
    MIN_TEXT_LENGTH = 50  # Minimum characters required
    MAX_CONCURRENT_UPLOADS = 100  # Maximum documents in processing queue
    
    ALLOWED_MIME_TYPES: Set[str] = {
        'application/pdf',
        'text/plain',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/html',
        'text/markdown'
    }
    
    ALLOWED_EXTENSIONS: Set[str] = {
        '.pdf', '.txt', '.doc', '.docx', '.html', '.md'
    }
    
    # Document type mapping
    MIME_TO_DOCTYPE = {
        'application/pdf': 'contract',
        'application/msword': 'contract',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'contract',
        'text/plain': 'general',
        'text/html': 'report',
        'text/markdown': 'documentation'
    }

# --- Router Setup ---

router = APIRouter(
    prefix="/documents",
    tags=["Documents"],
    responses={
        404: {"description": "Not found"},
        413: {"description": "File too large"},
        415: {"description": "Unsupported media type"},
        422: {"description": "Validation error"}
    }
)

logger = logging.getLogger(__name__)

# --- In-memory status tracking (use Redis in production) ---
processing_status: Dict[str, Dict] = {}
status_lock = asyncio.Lock()

# Thread pool for CPU-intensive tasks
executor = ThreadPoolExecutor(max_workers=2)

# --- Response Models ---

class ErrorDetail(BaseModel):
    code: str
    message: str
    field: Optional[str] = None

class ErrorResponse(BaseModel):
    error: str
    details: List[ErrorDetail]
    tracking_id: Optional[str] = None
    timestamp: datetime = Field(default_factory=datetime.utcnow)

class UploadResponse(BaseModel):
    tracking_id: str
    message: str
    status_url: str
    webhook_url: Optional[str] = None

class ProcessingStatusResponse(BaseModel):
    tracking_id: str
    status: str
    filename: str
    file_size: int
    created_at: str
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    progress: int
    stages: Dict[str, Dict[str, Any]]
    error: Optional[str] = None
    estimated_completion: Optional[float] = None

# --- Validation Functions ---

async def validate_file(file: UploadFile) -> None:
    """Comprehensive file validation with security checks."""
    # Check file size first (before reading)
    if file.size and file.size > DocumentConfig.MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Maximum size: {DocumentConfig.MAX_FILE_SIZE / 1024 / 1024:.1f}MB"
        )
    
    # Validate extension
    if file.filename:
        ext = Path(file.filename).suffix.lower()
        if ext not in DocumentConfig.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail=f"File type '{ext}' not supported. Allowed types: {', '.join(DocumentConfig.ALLOWED_EXTENSIONS)}"
            )
    
    # Validate MIME type from headers
    if file.content_type not in DocumentConfig.ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Content type '{file.content_type}' not supported"
        )
    
    # Read first chunk for magic byte validation
    first_chunk = await file.read(8192)
    await file.seek(0)  # Reset for later reading
    
    # Verify actual content type with magic bytes
    try:
        actual_mime = magic.from_buffer(first_chunk, mime=True)
        if actual_mime not in DocumentConfig.ALLOWED_MIME_TYPES:
            # Allow text/x-* variants for text files
            if not (actual_mime.startswith('text/') and 'text/plain' in DocumentConfig.ALLOWED_MIME_TYPES):
                raise HTTPException(
                    status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                    detail=f"File content detected as '{actual_mime}', which is not allowed"
                )
    except Exception as e:
        logger.error(f"Magic byte detection failed: {e}")
        # If magic detection fails, reject the file for security
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Unable to verify file content type. File rejected for security reasons."
        )

# --- Text Extraction ---

async def extract_pdf_text_cpu_bound(contents: bytes, metadata: Dict[str, Any]) -> str:
    """CPU-bound PDF extraction to run in thread pool."""
    return await extract_pdf_text(contents, metadata)

async def extract_docx_text_cpu_bound(contents: bytes, metadata: Dict[str, Any]) -> str:
    """CPU-bound DOCX extraction to run in thread pool."""
    return await extract_docx_text(contents, metadata)

async def extract_text_from_file(file: UploadFile, use_streaming: bool = False) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from various file formats with proper encoding detection.
    
    Args:
        file: The uploaded file
        use_streaming: Whether to use streaming for large files
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    metadata = {
        "original_filename": file.filename,
        "mime_type": file.content_type,
        "file_size": file.size
    }
    
    try:
        # Handle large files with streaming
        if use_streaming and file.size > DocumentConfig.LARGE_FILE_THRESHOLD:
            return await extract_text_streaming(file, metadata)
        
        # Read entire file for small files
        contents = await file.read()
        
        # Route to appropriate extractor
        if file.content_type == 'application/pdf':
            if not HAS_PDF:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="PDF support not installed. Please contact administrator."
                )
            # Run CPU-bound PDF extraction in thread pool
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(executor, extract_pdf_text, contents, metadata)
            
        elif file.content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            if not HAS_DOCX:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Word document support not installed."
                )
            # Run CPU-bound DOCX extraction in thread pool
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(executor, extract_docx_text, contents, metadata)
            
        else:
            # Plain text with encoding detection
            text = await extract_plain_text(contents, metadata)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Text extraction failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to extract text from {file.content_type}: {str(e)}"
        )
    
    # Validate extracted text
    if not text or len(text.strip()) < DocumentConfig.MIN_TEXT_LENGTH:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Document contains insufficient text content (minimum {DocumentConfig.MIN_TEXT_LENGTH} characters required)"
        )
    
    metadata["text_length"] = len(text)
    metadata["text_hash"] = hashlib.sha256(text.encode()).hexdigest()[:16]
    
    return text, metadata

async def extract_plain_text(contents: bytes, metadata: Dict[str, Any]) -> str:
    """Extract text with robust encoding detection."""
    # Try chardet for encoding detection
    detected = chardet.detect(contents)
    encoding = detected.get('encoding', 'utf-8')
    confidence = detected.get('confidence', 0)
    
    metadata["detected_encoding"] = encoding
    metadata["encoding_confidence"] = confidence
    
    # Try primary encoding
    try:
        if confidence > 0.8:
            return contents.decode(encoding)
    except (UnicodeDecodeError, TypeError):
        pass
    
    # Fallback encodings
    for enc in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
        try:
            return contents.decode(enc)
        except UnicodeDecodeError:
            continue
    
    # Last resort: decode with replacement
    return contents.decode('utf-8', errors='replace')

async def extract_pdf_text(contents: bytes, metadata: Dict[str, Any]) -> str:
    """Extract text from PDF with metadata."""
    # Try PyMuPDF first if available (5-10x faster)
    try:
        import fitz  # PyMuPDF
        pdf_doc = fitz.open(stream=contents, filetype="pdf")
        text_parts = []
        metadata["page_count"] = len(pdf_doc)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            page_text = page.get_text()
            if page_text:
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
        
        pdf_doc.close()
        logger.info("Used PyMuPDF for fast PDF extraction")
        return '\n\n'.join(text_parts)
        
    except ImportError:
        # Fallback to PyPDF2
        logger.info("PyMuPDF not available, using PyPDF2")
        pdf_file = io.BytesIO(contents)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        metadata["page_count"] = len(reader.pages)
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            except Exception as e:
                logger.warning(f"Failed to extract page {page_num}: {e}")
        
        return '\n\n'.join(text_parts)

async def extract_docx_text(contents: bytes, metadata: Dict[str, Any]) -> str:
    """Extract text from Word documents."""
    doc_file = io.BytesIO(contents)
    doc = docx.Document(doc_file)
    
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    metadata["paragraph_count"] = len(doc.paragraphs)
    metadata["table_count"] = len(doc.tables)
    
    return '\n\n'.join(paragraphs)

async def extract_text_streaming(file: UploadFile, metadata: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
    """Stream large file processing to avoid memory issues."""
    # For large files, save to temp file first
    temp_path = f"/tmp/kairos_upload_{uuid.uuid4()}"
    
    try:
        async with aiofiles.open(temp_path, 'wb') as temp_file:
            chunk_size = 1024 * 1024  # 1MB chunks
            while chunk := await file.read(chunk_size):
                await temp_file.write(chunk)
        
        # Process from temp file
        async with aiofiles.open(temp_path, 'rb') as temp_file:
            contents = await temp_file.read()
            
        # Route to appropriate processor
        if file.content_type == 'application/pdf':
            text = await extract_pdf_text(contents, metadata)
        else:
            text = await extract_plain_text(contents, metadata)
            
        return text, metadata
        
    finally:
        # Cleanup temp file
        try:
            Path(temp_path).unlink()
        except:
            pass

# --- Document Processing ---

async def process_document_async(
    tracking_id: str,
    text_content: str,
    metadata: Dict[str, Any],
    vector_service: VectorService,
    graph_service: GraphService,
    llm_client: Any,
    webhook_url: Optional[str] = None
):
    """
    Asynchronous document processing with detailed progress tracking.
    """
    try:
        async with status_lock:
            processing_status[tracking_id]["status"] = "processing"
            processing_status[tracking_id]["started_at"] = datetime.utcnow().isoformat()
        
        # Stage 1: Vector Processing (50% of progress)
        async with status_lock:
            processing_status[tracking_id]["stages"]["vector"] = {
                "status": "processing",
                "started_at": datetime.utcnow().isoformat()
            }
        
        try:
            vector_result = await vector_service.add_document(
                text=text_content,
                metadata={
                    "document_id": tracking_id,
                    **metadata
                }
            )
            
            async with status_lock:
                processing_status[tracking_id]["stages"]["vector"] = {
                    "status": "completed" if vector_result.get("success") else "failed",
                    "chunks_created": vector_result.get("chunks_created", 0),
                    "processing_time": vector_result.get("processing_time", 0),
                    "completed_at": datetime.utcnow().isoformat()
                }
                processing_status[tracking_id]["progress"] = 50
                
        except Exception as e:
            logger.error(f"Vector processing failed for {tracking_id}: {e}")
            async with status_lock:
                processing_status[tracking_id]["stages"]["vector"] = {
                    "status": "failed",
                    "error": str(e),
                    "completed_at": datetime.utcnow().isoformat()
                }
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Vector processing failed: {str(e)}"
            )
        
        # Stage 2: Graph Extraction (remaining 50%)
        async with status_lock:
            processing_status[tracking_id]["stages"]["graph"] = {
                "status": "processing",
                "started_at": datetime.utcnow().isoformat()
            }
        
        try:
            graph_result = await graph_service.extract_and_build_graph(
                text=text_content,
                llm_client=llm_client,
                document_id=tracking_id,
                document_type=metadata.get("document_type", "general"),
                metadata=metadata
            )
            
            async with status_lock:
                processing_status[tracking_id]["stages"]["graph"] = {
                    "status": "completed" if graph_result else "failed",
                    "entities_extracted": len(graph_result.nodes) if graph_result else 0,
                    "relationships_found": len(graph_result.edges) if graph_result else 0,
                    "completed_at": datetime.utcnow().isoformat()
                }
                processing_status[tracking_id]["progress"] = 100
                
        except Exception as e:
            logger.error(f"Graph extraction failed for {tracking_id}: {e}")
            async with status_lock:
                processing_status[tracking_id]["stages"]["graph"] = {
                    "status": "failed",
                    "error": str(e),
                    "completed_at": datetime.utcnow().isoformat()
                }
                # Don't raise here - partial success is acceptable
        
        # Final status update
        async with status_lock:
            processing_status[tracking_id]["status"] = "completed"
            processing_status[tracking_id]["completed_at"] = datetime.utcnow().isoformat()
            
            # Calculate total processing time
            started = datetime.fromisoformat(processing_status[tracking_id]["started_at"])
            completed = datetime.utcnow()
            processing_status[tracking_id]["total_processing_time"] = (completed - started).total_seconds()
        
        # Send webhook if configured
        if webhook_url:
            await send_webhook_notification(webhook_url, {
                "tracking_id": tracking_id,
                "status": "completed",
                "stages": processing_status[tracking_id]["stages"]
            })
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Processing failed for {tracking_id}: {e}", exc_info=True)
        async with status_lock:
            processing_status[tracking_id]["status"] = "failed"
            processing_status[tracking_id]["error"] = str(e)
            processing_status[tracking_id]["completed_at"] = datetime.utcnow().isoformat()
        
        if webhook_url:
            await send_webhook_notification(webhook_url, {
                "tracking_id": tracking_id,
                "status": "failed",
                "error": str(e)
            })

async def send_webhook_notification(webhook_url: str, payload: Dict[str, Any]):
    """Send webhook notification for processing completion."""
    import httpx
    
    try:
        async with httpx.AsyncClient() as client:
            response = await client.post(
                webhook_url,
                json=payload,
                timeout=5.0,
                headers={"Content-Type": "application/json"}
            )
            response.raise_for_status()
    except Exception as e:
        logger.error(f"Webhook notification failed: {e}")

# --- API Endpoints ---

@router.post("/upload", 
    response_model=UploadResponse,
    status_code=status.HTTP_202_ACCEPTED,
    responses={
        202: {"description": "Document accepted for processing"},
        413: {"description": "File too large"},
        415: {"description": "Unsupported file type"},
        422: {"description": "Invalid document content"}
    }
)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(..., description="Document file (max 10MB)"),
    document_type: Optional[str] = Form(None, description="Document type (contract, policy, etc.)"),
    title: Optional[str] = Form(None, description="Document title"),
    description: Optional[str] = Form(None, description="Document description"),
    tags: Optional[str] = Form(None, description="Comma-separated tags"),
    webhook_url: Optional[str] = Form(None, description="Webhook URL for completion notification"),
    priority: str = Form("normal", description="Processing priority (low, normal, high)"),
    vector_service: VectorService = Depends(get_vector_service),
    graph_service: GraphService = Depends(get_graph_service),
    llm_client: Any = Depends(get_llm_client)
):
    """
    Upload and process a document asynchronously.
    
    Returns a tracking ID immediately and processes the document in the background.
    Use the tracking ID to check processing status.
    """
    # Check queue depth to prevent memory issues
    async with status_lock:
        processing_count = sum(1 for s in processing_status.values() 
                             if s["status"] in ["queued", "processing"])
        if processing_count >= DocumentConfig.MAX_CONCURRENT_UPLOADS:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=f"Server busy. Maximum {DocumentConfig.MAX_CONCURRENT_UPLOADS} documents in processing queue. Please try again later."
            )
    
    # Validate file first
    await validate_file(file)
    
    # Generate tracking ID
    tracking_id = f"doc_{uuid.uuid4()}"
    
    # Determine if we need streaming for large files
    use_streaming = file.size > DocumentConfig.LARGE_FILE_THRESHOLD
    
    # Extract text and metadata
    try:
        text_content, extraction_metadata = await extract_text_from_file(file, use_streaming)
    except HTTPException as e:
        # Re-raise with tracking ID for debugging
        raise HTTPException(
            status_code=e.status_code,
            detail=ErrorResponse(
                error="Text extraction failed",
                details=[ErrorDetail(code="EXTRACTION_FAILED", message=e.detail)],
                tracking_id=tracking_id
            ).dict()
        )
    
    # Auto-detect document type if not provided
    if not document_type:
        document_type = DocumentConfig.MIME_TO_DOCTYPE.get(file.content_type, "general")
    
    # Parse tags
    tag_list = [tag.strip() for tag in tags.split(",")] if tags else []
    
    # Sanitize filename to prevent directory traversal
    safe_filename = Path(file.filename).name if file.filename else f"unnamed_{uuid.uuid4()}"
    
    # Build complete metadata
    full_metadata = {
        "filename": safe_filename,  # Sanitized filename
        "original_filename": file.filename,  # Keep original for reference
        "file_size": file.size,
        "mime_type": file.content_type,
        "document_type": document_type,
        "title": title or safe_filename,
        "description": description,
        "tags": tag_list,
        "priority": priority,
        "uploaded_at": datetime.utcnow().isoformat(),
        **extraction_metadata
    }
    
    # Initialize processing status
    async with status_lock:
        processing_status[tracking_id] = {
            "tracking_id": tracking_id,
            "status": "queued",
            "filename": safe_filename,  # Use sanitized filename
            "file_size": file.size,
            "created_at": datetime.utcnow().isoformat(),
            "progress": 0,
            "stages": {
                "vector": {"status": "pending"},
                "graph": {"status": "pending"}
            },
            "metadata": full_metadata
        }
    
    # Queue background processing
    background_tasks.add_task(
        process_document_async,
        tracking_id=tracking_id,
        text_content=text_content,
        metadata=full_metadata,
        vector_service=vector_service,
        graph_service=graph_service,
        llm_client=llm_client,
        webhook_url=webhook_url
    )
    
    logger.info(f"Document {file.filename} queued for processing with ID {tracking_id}")
    
    return UploadResponse(
        tracking_id=tracking_id,
        message="Document queued for processing",
        status_url=f"/documents/status/{tracking_id}",
        webhook_url=webhook_url
    )

@router.get("/status/{tracking_id}", 
    response_model=ProcessingStatusResponse,
    responses={
        404: {"description": "Document not found"}
    }
)
async def get_processing_status(tracking_id: str):
    """
    Get the current processing status of a document.
    
    Returns detailed information about processing stages and progress.
    """
    async with status_lock:
        if tracking_id not in processing_status:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found or processing not started"
            )
        
        status = processing_status[tracking_id].copy()
    
    # Add time estimates for in-progress documents
    if status["status"] == "processing" and "started_at" in status:
        started = datetime.fromisoformat(status["started_at"])
        elapsed = (datetime.utcnow() - started).total_seconds()
        
        # Estimate based on progress
        if status["progress"] > 0:
            estimated_total = elapsed / (status["progress"] / 100)
            status["estimated_completion"] = max(0, estimated_total - elapsed)
        else:
            status["estimated_completion"] = 30  # Default estimate
    
    return ProcessingStatusResponse(**status)

@router.get("/status/{tracking_id}/stream")
async def stream_processing_status(tracking_id: str, updates_per_sec: int = 2):
    """
    Stream processing status updates via Server-Sent Events.
    
    Provides real-time updates as the document is processed.
    
    Args:
        tracking_id: Document tracking ID
        updates_per_sec: Maximum updates per second (1-10, default 2)
    """
    # Validate update rate
    updates_per_sec = max(1, min(10, updates_per_sec))
    update_interval = 1.0 / updates_per_sec
    
    async with status_lock:
        if tracking_id not in processing_status:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
    
    async def event_generator():
        last_update = None
        
        while True:
            async with status_lock:
                status = processing_status.get(tracking_id)
            
            if status:
                # Only send if there's an update
                current_update = {
                    "status": status["status"],
                    "progress": status["progress"],
                    "stages": status["stages"]
                }
                
                if current_update != last_update:
                    yield {
                        "event": "status",
                        "data": json.dumps(status)
                    }
                    last_update = current_update
                
                # End stream when processing is complete
                if status["status"] in ["completed", "failed"]:
                    yield {
                        "event": "complete",
                        "data": json.dumps({"tracking_id": tracking_id})
                    }
                    break
            
            await asyncio.sleep(update_interval)
    
    return EventSourceResponse(event_generator())

@router.post("/upload/batch",
    response_model=Dict[str, Any],
    status_code=status.HTTP_202_ACCEPTED
)
async def upload_batch(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(..., description="Multiple files (max 10)"),
    document_type: Optional[str] = Form(None),
    priority: str = Form("normal"),
    webhook_url: Optional[str] = Form(None),
    vector_service: VectorService = Depends(get_vector_service),
    graph_service: GraphService = Depends(get_graph_service),
    llm_client: Any = Depends(get_llm_client)
):
    """
    Upload multiple documents for batch processing.
    
    Maximum 10 files per batch. Each file is processed independently.
    """
    if len(files) > 10:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Maximum 10 files allowed per batch"
        )
    
    batch_id = f"batch_{uuid.uuid4()}"
    results = []
    
    for file in files:
        try:
            # Validate each file
            await validate_file(file)
            
            # Generate tracking ID
            tracking_id = f"doc_{uuid.uuid4()}"
            
            # Extract text
            text_content, metadata = await extract_text_from_file(file)
            
            # Initialize status
            async with status_lock:
                processing_status[tracking_id] = {
                    "tracking_id": tracking_id,
                    "status": "queued",
                    "filename": file.filename,
                    "file_size": file.size,
                    "created_at": datetime.utcnow().isoformat(),
                    "progress": 0,
                    "stages": {
                        "vector": {"status": "pending"},
                        "graph": {"status": "pending"}
                    },
                    "batch_id": batch_id,
                    "metadata": metadata
                }
            
            # Queue for processing
            background_tasks.add_task(
                process_document_async,
                tracking_id=tracking_id,
                text_content=text_content,
                metadata={
                    **metadata,
                    "document_type": document_type or "general",
                    "batch_id": batch_id,
                    "priority": priority
                },
                vector_service=vector_service,
                graph_service=graph_service,
                llm_client=llm_client,
                webhook_url=webhook_url
            )
            
            results.append({
                "filename": file.filename,
                "tracking_id": tracking_id,
                "status": "queued"
            })
            
        except HTTPException as e:
            results.append({
                "filename": file.filename,
                "error": e.detail,
                "status": "rejected"
            })
            logger.warning(f"Batch upload: rejected {file.filename}: {e.detail}")
    
    successful = [r for r in results if r["status"] == "queued"]
    
    return {
        "batch_id": batch_id,
        "total_files": len(files),
        "accepted": len(successful),
        "rejected": len(files) - len(successful),
        "documents": results,
        "batch_status_url": f"/documents/batch/{batch_id}/status"
    }

@router.get("/batch/{batch_id}/status")
async def get_batch_status(batch_id: str):
    """Get status of all documents in a batch."""
    async with status_lock:
        batch_docs = [
            status for tracking_id, status in processing_status.items()
            if status.get("metadata", {}).get("batch_id") == batch_id
        ]
    
    if not batch_docs:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Batch not found"
        )
    
    summary = {
        "batch_id": batch_id,
        "total": len(batch_docs),
        "completed": len([d for d in batch_docs if d["status"] == "completed"]),
        "failed": len([d for d in batch_docs if d["status"] == "failed"]),
        "processing": len([d for d in batch_docs if d["status"] == "processing"]),
        "queued": len([d for d in batch_docs if d["status"] == "queued"])
    }
    
    return {
        "summary": summary,
        "documents": [
            {
                "tracking_id": doc["tracking_id"],
                "filename": doc.get("filename", "unknown"),  # Safe get
                "status": doc["status"],
                "progress": doc["progress"]
            }
            for doc in batch_docs
        ]
    }

# --- Cleanup Task ---

async def cleanup_old_statuses():
    """Remove status entries older than 24 hours."""
    cutoff = datetime.utcnow() - timedelta(hours=24)
    
    async with status_lock:
        to_remove = []
        for tracking_id, status in processing_status.items():
            created = datetime.fromisoformat(status["created_at"])
            if created < cutoff and status["status"] in ["completed", "failed"]:
                to_remove.append(tracking_id)
        
        for tracking_id in to_remove:
            del processing_status[tracking_id]
        
        if to_remove:
            logger.info(f"Cleaned up {len(to_remove)} old status entries")

# Schedule cleanup to run periodically (implement with APScheduler or similar)
# scheduler.add_job(cleanup_old_statuses, 'interval', hours=1)